from selenium import webdriver
from selenium.common.exceptions import NoSuchElementException

import time
import random

from bs4 import BeautifulSoup as bs
from docx import Document

from datetime import datetime as dt


class Ul(object):
    """docstring for Ul."""
    def __init__(self, object, name=None, address=None, info=None):
        self.innul = object[0]
        self.ogrn = object[1]
        self.name = name
        self.address = address
        self.info = info


    def bik(self, driver, id_name):
        bikn = driver.find_element_by_id(id_name)
        bik = [x for x in '000000000']
        for i in bik:
            time.sleep(.05)
            bikn.send_keys(i)
        return True

    def button(self, driver, id_name):
        try:
            driver.find_element_by_id(id_name).click()
            time.sleep(random.randint(2, 7))
            return True
        except NoSuchElementException:
            driver.find_element_by_class_name(id_name).click()
            time.sleep(random.randint(2, 7))
            return True

    def input_key(self, driver, form_id, key):
        username = driver.find_element_by_id(form_id)
        for i in key:
            time.sleep(.05)
            username.send_keys(i)
        return username

    def radio_click(self, driver, id):
        driver.find_element_by_id(id).click()
        time.sleep(random.randint(2, 7))
        return True

    def get_text(self, driver, id):
        time.sleep(random.randint(2, 7))
        try:
            # text = driver.find_element_by_id(id).text
            html = driver.find_element_by_id(id).get_attribute('innerHTML')
            return html
        except NoSuchElementException:
            html = driver.find_element_by_class_name(id).get_attribute('innerHTML')
            # rows = driver.find_elements_by_id(id)
            # text = [row.text for row in rows]
            return html

    def cap(self, driver):
        captcha = driver.find_element_by_id('captcha')
        f = [x for x in input('введите капчу: ')]
        for i in f:
            time.sleep(.05)
            captcha.send_keys(i)
        return True

    def cap_loop(self, username, driver, btn_id=None):
        er = True
        while er:
            self.cap(driver)
            if btn_id:
                self.button(driver, btn_id)
            else:
                username.submit()
            try:
                er = driver.find_element_by_id('errors_captcha').text
                print(er)
            except NoSuchElementException:
                er = False
        print('Цифры с картинки введены верно')

    def services(self):
        dict_div = {}
        key = [x for x in self.innul]
        if self.ogrn:
            ogrn = [x for x in self.ogrn]
        # driver = webdriver.Firefox(executable_path='C:\\Python\geckodriver.exe')
        driver = webdriver.Chrome()

        # url = 'https://fedresurs.ru'
        url = services[0]
        driver.get(url)
        username = driver.find_element_by_css_selector("input[name='searchString']")
        for i in key:
            time.sleep(.05)
            username.send_keys(i)
        driver.find_element_by_class_name("btn").click()
        dict_div[url] = self.get_text(driver, 'tab-content')
        # fedresurs = self.get_text(driver, 'search-result')

        if self.ogrn:
            url = services[1]
            # 'https://service.nalog.ru/uwsfind.do'
            driver.get(url)
            username = self.input_key(driver, 'ogrnUl', ogrn)
            self.cap_loop(username, driver, 'btnSearch')
            table = self.get_text(driver, 'tableResultData')
            if table:
                dict_div[url] = table
            else:
                dict_div[url] = self.get_text(driver, 'pnlResult')
        else:
            print('не проверяли')

        url = services[2]
        # 'https://service.nalog.ru/disqualified.do'
        driver.get(url)
        username = self.input_key(driver, 'orgInn', key)
        self.cap_loop(username, driver, 'float-right')
        # third = self.get_text(driver, 'resultPanel')
        dict_div[url] = self.get_text(driver, 'resultPanel')


        url = services[3]
        # 'http://zakupki.gov.ru/epz/dishonestsupplier/quicksearch/search.html'
        driver.get(url)
        username = self.input_key(driver, 'searchString', key)
        username.submit()
        time.sleep(random.randint(2, 7))
        try:
            # fourth = driver.find_element_by_css_selector('table.searchDocTable').text
            dict_div[url] = self.get_text(driver, 'margBtm10')
        except NoSuchElementException:
            dict_div[url] = 'Поиск не дал результатов.'


        url = services[4]
        # 'https://service.nalog.ru/svl.do'
        driver.get(url)
        username = self.input_key(driver, 'svlform_inn', key)
        self.cap_loop(username, driver, 'btn-ok')
        try:
            # fifth = driver.find_element_by_class_name('container').text
            dict_div[url] = self.get_text(driver, 'container')
        except NoSuchElementException:
            # fifth = driver.find_element_by_class_name('panel').text
            dict_div[url] = self.get_text(driver, 'panel')

        url = services[5]
        # 'http://bankrot.fedresurs.ru/DebtorsSearch.aspx'
        driver.get(url)
        username = self.input_key(driver, 'ctl00_cphBody_OrganizationCode1_CodeTextBox', key)
        self.button(driver, 'ctl00_cphBody_btnSearch')
        # sixth = self.get_text(driver, 'ctl00_cphBody_upList')
        dict_div[url] = self.get_text(driver, 'ctl00_cphBody_upList')


        url = services[6]
        # 'https://service.nalog.ru/zd.do'
        driver.get(url)
        username = self.input_key(driver, 'inn', key)
        captcha = driver.find_element_by_id('captcha')
        self.cap_loop(username, driver, 'btn_send')
        # seventh = self.get_text(driver, 'pnlResults')
        dict_div[url] = self.get_text(driver, 'pnlResults')

        url = services[7]
        # 'https://service.nalog.ru/bi.do'
        driver.get(url)
        self.radio_click(driver, 'unirad_0')
        username = self.input_key(driver, 'innPRS', key)
        self.bik(driver, 'bikPRS')
        self.cap_loop(username, driver, 'btnSearch')
        # eighth = self.get_text(driver, 'pnlResultData')
        dict_div[url] = self.get_text(driver, 'pnlResultData')

        # self.info = (
        #         ('https://fedresurs.ru', fedresurs),
        #         ('https://service.nalog.ru/uwsfind.do', second),
        #         ('https://service.nalog.ru/disqualified.do', third),
        #         ('http://zakupki.gov.ru/epz/dishonestsupplier/quicksearch/search.html', fourth),
        #         ('https://service.nalog.ru/svl.do', fifth),
        #         ('http://bankrot.fedresurs.ru/DebtorsSearch.aspx', sixth),
        #         ('https://service.nalog.ru/zd.do', seventh),
        #         ('https://service.nalog.ru/bi.do', eighth)
        #             )
        # print(input('close?: ', ))
        self.info = dict_div

        driver.close()

if __name__ == '__main__':
    services = (
        'https://fedresurs.ru',
        'https://service.nalog.ru/uwsfind.do',
        'https://service.nalog.ru/disqualified.do',
        'http://zakupki.gov.ru/epz/dishonestsupplier/quicksearch/search.html',
        'https://service.nalog.ru/svl.do',
        'http://bankrot.fedresurs.ru/DebtorsSearch.aspx',
        'https://service.nalog.ru/zd.do',
        'https://service.nalog.ru/bi.do',
    )
    headers = (
        'Федресурс',
        'Сведения о юридических лицах и индивидуальных предпринимателях, в отношении которых представлены документы для государственной регистрации',
        'Поиск сведений в реестре дисквалифицированных лиц',
        'Портал Закупок Результаты поиска',
        'Сведения о лицах, в отношении которых факт невозможности участия (осуществления руководства) в организации установлен (подтвержден) в судебном порядке',
        'Единый федеральный реестр сведений о банкротстве',
        'Сведения о юридических лицах, имеющих задолженность по уплате налогов и/или не представляющих налоговую отчетность более года',
        'Система информирования банков о состоянии обработки электронных документов (311-П, 440-П)',
    )



    key = ''
    ogrn = ''

    file = ('{} - {}{}').format(key, dt.today().strftime("%d-%m-%Y"), '.docx')

    key = [x for x in key]
    ogrn = [x for x in ogrn]

    ul = Ul((key, ogrn),)
    ul.services()

    dict_div = ul.info

    document = Document()
    index_header = 0
    for k, v in dict_div.items():
        soup = bs(v, 'lxml')


        head = headers[index_header]
        index_header += 1

        document.add_heading(head, level=2)
        try:
            # soup_table = soup.find('table', class_='search-result')
            soup_table = soup.find('table')
            print('soup_table is ok!')
        except Exception as e:
            print('soup_table is not ok!' + e)

        if soup_table:
            soup_rows = soup_table.find_all('tr')
            try:
                soup_cols = soup_table.find_all('tr')[1].find_all('td')
            except IndexError:
                soup_cols = soup_table.find_all('tr')[0].find_all('td')

            len_rows = len(soup_rows)
            len_cols = len(soup_cols)

            doc_table = document.add_table(rows=len_rows, cols=len_cols, style='Table Grid')

            for idx, soup_row in enumerate(soup_rows):
                soup_cols = soup_row.find_all('td') or soup_row.find_all('th')
                # print('idx -', idx, 'cols -', soup_cols, 'rows - ', soup_rows)
                # print('idx', idx)
                doc_cells = doc_table.rows[idx].cells
                # print(doc_cells, type(doc_cells))
                for i, soup_col in enumerate(soup_cols):
                    # print('i', i)
                    doc_cells[i].text = soup_col.text
        else:
            try:
                nf = soup.find('p').text
            except NoSuchElementException:
                nf = soup.get_text()
            except AttributeError:
                nf = soup.get_text()

            document.add_paragraph(nf)
            # print('ой, тут не таблица')
        document.add_paragraph()

    document.save(file)
    print('сохраняем документ')
