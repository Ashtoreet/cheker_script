from service import Service
from chek_obj import CheckObj
import datetime
from docx import Document
from selenium import webdriver
from selenium.common.exceptions import NoSuchElementException
from bs4 import BeautifulSoup as bs
import sys


services = (
    'https://fedresurs.ru',
    'https://service.nalog.ru/uwsfind.do',
    'https://service.nalog.ru/disqualified.do',
    'http://zakupki.gov.ru/epz/dishonestsupplier/quicksearch/search.html',
    'https://service.nalog.ru/svl.do',
    'http://bankrot.fedresurs.ru/DebtorsSearch.aspx',
    'https://service.nalog.ru/zd.do',
    'https://service.nalog.ru/bi.do',
)
headers = (
        'Федресурс',
        'Сведения о юридических лицах и индивидуальных предпринимателях, '
        'в отношении которых представлены документы для государственной регистрации',
        'Поиск сведений в реестре дисквалифицированных лиц',
        'Портал Закупок Результаты поиска',
        'Сведения о лицах, в отношении которых факт невозможности участия '
        '(осуществления руководства) в организации установлен (подтвержден) в судебном порядке',
        'Единый федеральный реестр сведений о банкротстве. Должники.',
        'Сведения о юридических лицах, имеющих задолженность по уплате налогов и/или'
        ' не представляющих налоговую отчетность более года',
        'Система информирования банков о состоянии обработки электронных документов (311-П, 440-П)',
        'Единый федеральный реестр сведений о банкротстве. Дисквалифицированные лица.',
    )


def add_infotable(info):
    table = document.add_table(rows=1, cols=2, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Источник'
    hdr_cells[1].text = 'Результат'

    for k, v in info:
        row_cells = table.add_row().cells
        row_cells[0].text = k
        if v:
            row_cells[1].text = v
        else:
            row_cells[1].text = 'Пусто'


# data acquisition
def user_input():
    var = False
    print('Программа заработала, для выхода нажмите Ctrl + C.')

    while var is False:
        key = input('Введите ключ: ', ).strip()
        try:
            int(key)

            if len(key) is 10 or len(key) is 13:
                checkobj = CheckObj([x for x in key])
                fs = checkobj.ruprofile()
                print('ul = True')
                return(checkobj, fs)

            elif len(key) is 12 or len(key) is 15:
                checkobj = CheckObj([x for x in key])

                fs = checkobj.ruprofile()
                return (checkobj, fs)

            else:
                print('Неверные данные.')
                var = False
        except ValueError:
            if not key:
                return False
            print('Неверные данные.')
            var = False


dict_data = []
checkobj, fs = user_input()

# checkobj.ruprofile()

if fs is 1 or fs is 3:
    ul = False
elif fs is 0:
    sys.exit()
else:
    ul = True

driver = webdriver.Chrome()
# driver = webdriver.Firefox(executable_path='C:\\Python\geckodriver.exe')

if fs is 1:
    print("Это физик")
    service = Service(checkobj.inn, checkobj.ogrn, ul)
    service.fiz(checkobj.boss_name, driver)
    dict_data.append(service.dict_service)
else:
    print('Это не физик')

    # ogrn = [x for x in checkobj.ogrn]
    # list_key = [x for x in checkobj.key]
    # print(checkobj.bosses_inn)
    """
    доработать
    for i in checkobj.bosses_inn:
        print(i)
        checkobj_f_boss = CheckObj(key)
        ul = checkobj_f_boss.egrul()
        service_f_boss = Service(i, ul)

        checkobj_f_boss.zachestnyibiznes()
        i = [x for x in i]
        service_f_boss.fiz(checkobj_f_boss.boss_name, i)
        dict_data.append(service_f_boss.dict_service)
        print(input('продолжить?', ))
    """

    if checkobj.date_result.days < 60:

        print('меньше 60 дней')
        if fs is 2:
            service = Service(checkobj.inn, checkobj.ogrn, ul)
            service.ur_min(driver)
        else:
            service = Service(checkobj.inn, checkobj.ogrn, ul)
            service.ip_min(driver)
    else:
        print('больше 60 дней')
        if fs is 2:
            service = Service(checkobj.inn, checkobj.ogrn, ul)
            service.ur_min(driver)
            service.ur_max(checkobj.boss_name, driver)

        else:
            service = Service(checkobj.inn, checkobj.ogrn, ul)
            service.ip_min(driver)
            service.ip_max(checkobj.boss_name, driver)

    dict_data.append(service.dict_service)
driver.close()

print('начинаем собирать документ')
document = Document()
today = datetime.date.today().strftime('%d.%m.%y')
key_inn = ''.join(checkobj.key)
filename = '{} - {}.docx'.format(key_inn, today)


if fs is 1 or fs is 3:
    document.add_heading('{}'.format(checkobj.boss_name, level=1))
    document.add_heading('ИНН: {}'.format(key_inn), level=2)
    # filename = '{} - {}.docx'.format(checkobj.key, today)
else:
    document.add_heading('{}'.format(checkobj.company_name), level=1)
    document.add_heading('ИНН: {}'.format(key_inn), level=2)
    document.add_heading('ОГРН: {}'.format(''.join(checkobj.ogrn)), level=2)
    document.add_heading('{}'.format(checkobj.boss_name), level=2)

for data in dict_data:
    # index_header = 1
    for k, v in service.dict_service.items():
        soup = bs(v, 'lxml')
        # headers[services.index('https://fedresurs.ru')]
        head = headers[services.index(k)]
        # head = headers[index_header]
        # index_header += 1

        document.add_heading(head, level=2)
        try:
            # soup_table = soup.find('table', class_='search-result')
            soup_tables = soup.find_all('table')
            print('soup_table is ok!')
        except Exception as e:
            print('soup_table is not ok!' + e)

        if soup_tables:
            for soup_table in soup_tables:
                soup_rows = soup_table.find_all('tr')
                try:
                    soup_cols = soup_table.find_all('tr')[1].find_all('td')
                except IndexError:
                    soup_cols = soup_table.find_all('tr')[0].find_all('td')

                len_rows = len(soup_rows)
                len_cols = len(soup_cols)

                doc_table = document.add_table(rows=len_rows, cols=len_cols, style='Table Grid')

                for idx, soup_row in enumerate(soup_rows):
                    soup_cols = soup_row.find_all('td') or soup_row.find_all('th')
                    # print('idx -', idx, 'cols -', soup_cols, 'rows - ', soup_rows)
                    # print('idx', idx)
                    doc_cells = doc_table.rows[idx].cells
                    # print(doc_cells, type(doc_cells))
                    for i, soup_col in enumerate(soup_cols):
                        # print('i', i)
                        doc_cells[i].text = soup_col.text
        else:
            try:
                nf = soup.find('p').text
            except NoSuchElementException:
                nf = soup.get_text()
            except AttributeError:
                nf = soup.get_text()

            document.add_paragraph(nf)
            # print('ой, тут не таблица')
        document.add_paragraph()

document.add_page_break()
document.save(filename)


print('сохраняем документ')
