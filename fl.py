from selenium import webdriver
from selenium.common.exceptions import NoSuchElementException

import time
import random

from bs4 import BeautifulSoup as bs
from docx import Document

from datetime import datetime as dt


class Fl(object):
    """docstring for Fl."""
    def __init__(self, object, fio=None, info=None):
        self.innfl = object[0]
        self.ogrnip = object[1]
        self.fio = fio
        self.info = info

    def bik(self, driver, id_name):
        bikn = driver.find_element_by_id(id_name)
        bik = [x for x in '000000000']
        for i in bik:
            time.sleep(.05)
            bikn.send_keys(i)
        return True

    def button(self, driver, id_name):
        try:
            driver.find_element_by_id(id_name).click()
            time.sleep(random.randint(2, 7))
            return True
        except NoSuchElementException:
            driver.find_element_by_class_name(id_name).click()
            time.sleep(random.randint(2, 7))
            return True

    def input_key(self, driver, form_id, key):
        username = driver.find_element_by_id(form_id)
        for i in key:
            time.sleep(.05)
            username.send_keys(i)
        return username

    def radio_click(self, driver, id):
        driver.find_element_by_id(id).click()
        time.sleep(random.randint(2, 7))
        return True

    def get_text(self, driver, id):
        time.sleep(random.randint(2, 7))
        try:
            # text = driver.find_element_by_id(id).text
            html = driver.find_element_by_id(id).get_attribute('innerHTML')
            return html
        except NoSuchElementException:
            html = driver.find_element_by_class_name(id).get_attribute('innerHTML')
            # rows = driver.find_elements_by_id(id)
            # text = [row.text for row in rows]
            return html

    def cap(self, driver):
        captcha = driver.find_element_by_id('captcha')
        f = [x for x in input('введите капчу: ')]
        for i in f:
            time.sleep(.05)
            captcha.send_keys(i)
        return True

    def cap_loop(self, username, driver, btn_id=None):
        er = True
        while er:
            self.cap(driver)
            if btn_id:
                self.button(driver, btn_id)
            else:
                username.submit()
            try:
                er = driver.find_element_by_id('errors_captcha').text
                print(er)
            except NoSuchElementException:
                er = False
        print('Цифры с картинки введены верно')


    def services(self):
        dict_div = {}
        key = [x for x in self.innfl]
        if self.ogrnip:
            ogrnip = [x for x in self.ogrnip]
        driver = webdriver.Chrome()
        # driver = webdriver.Firefox(executable_path='C:\\Python\geckodriver.exe')



        # 1
        url = services[0]
        driver.get(url)
        # driver.get('http://zakupki.gov.ru/epz/dishonestsupplier/quicksearch/search.html')
        username = self.input_key(driver, 'searchString', key)
        username.submit()
        time.sleep(random.randint(2, 7))
        # try:
        dict_div[url] = self.get_text(driver, 'margBtm10')
            # first = driver.find_element_by_css_selector('table.searchDocTable').text
        # except NoSuchElementException:
        #     dict_div[url] = 'Ничего не найдено.'

        # 2
        url = services[1]
        driver.get(url)
        # driver.get('https://service.nalog.ru/bi.do')
        self.radio_click(driver, 'unirad_0')
        username = self.input_key(driver, "innPRS", key)
        self.bik(driver, "bikPRS")
        self.cap_loop(username, driver, 'btnSearch')
        dict_div[url] = self.get_text(driver, 'pnlResultData')

        # 3
        url = services[2]
        driver.get(url)
        time.sleep(random.randint(2, 7))
        # driver.get('http://se.fedresurs.ru/IndividualEntrepreneurs')
        self.input_key(driver, 'ctl00_MainContent_txtCode', key)
        self.button(driver, 'ctl00_MainContent_btnSearch')
        time.sleep(random.randint(2, 7))
        # try:
        link = driver.find_element_by_class_name('fn')
        self.fio = link.text
        self.ogrnip = driver.find_element_by_id(
                'ctl00_MainContent_upnIndEntrepreneurList'
            ).text.split(' ').pop().replace(')', '')
        link.click()
        time.sleep(random.randint(2, 7))

        dict_div[url] = self.get_text(driver, 'vcard')
        print(self.fio)
        # except NoSuchElementException:
        #     third = driver.find_element_by_id('ctl00_MainContent_upnIndEntrepreneurList').text
        #     self.ogrnip = driver.find_element_by_id(
        #             'ctl00_MainContent_upnIndEntrepreneurList'
        #         ).text.split(' ').pop().replace(')', '')
        #     self.fio = None

        # # 4
        # driver.get('http://bankrot.fedresurs.ru/DebtorsSearch.aspx')
        # time.sleep(2)
        # self.radio_click(driver, 'ctl00_cphBody_rblDebtorType_1')
        # time.sleep(3)
        # self.input_key(driver, 'ctl00_cphBody_PersonCode1_CodeTextBox', key)
        # self.button(driver, 'ctl00_cphBody_btnSearch')
        # fourth = self.get_text(driver, 'ctl00_cphBody_upList')
        #
        # # 6
        # if self.ogrnip:
        #     driver.get('https://service.nalog.ru/uwsfind.do')
        #     self.radio_click(driver, 'unirad_1')
        #     username = self.input_key(driver, 'ogrnIp', self.ogrnip)
        #     self.cap_loop(username, driver, 'btnSearch')
        #     if self.fio is None:
        #         self.fio = driver.find_element_by_class_name('uws-result-item-value').text
        #         print(self.fio)
        #     sixth = self.get_text(driver, 'pnlResult')
        # else:
        #     sixth = 'Не проверяли'
        #
        # # 5, работает только с третьим пунктом(берет там name)
        # if self.fio:
        #     driver.get('http://bankrot.fedresurs.ru/DisqualificantsList.aspx')
        #     self.input_key(driver, 'ctl00_cphBody_tbFio', self.fio)
        #     self.button(driver, 'ctl00_cphBody_btnSearch')
        #     fifth = self.get_text(driver, 'ctl00_cphBody_upDisqList')
        # else:
        #     fifth = 'Имя не определено'
        #
        #
        # # 7, работает только с третьим пунктом(берет там name)
        # if self.fio:
        #     driver.get('https://service.nalog.ru/disqualified.do')
        #     userfam = self.fio.split()[0]
        #     print(userfam)
        #     usernam = self.fio.split()[1]
        #     print(usernam)
        #     userotch = self.fio.split()[2]
        #     print(userotch)
        #     self.input_key(driver, 'otch', userotch)
        #     self.input_key(driver, 'fam', userfam)
        #     username = self.input_key(driver, 'nam', usernam)
        #     self.cap_loop(username, driver, 'btn-ok')
        #     seventh = self.get_text(driver, 'resultPanel')
        # else:
        #     seventh = 'Имя не определено'
        #
        # # self.info = (
        # #             ('http://zakupki.gov.ru/epz/dishonestsupplier/quicksearch/search.html : ', first),
        # #             ('https://service.nalog.ru/bi.do', second),
        # #             ('https://service.nalog.ru/disqualified.do', third),
        # #             ('http://bankrot.fedresurs.ru/DebtorsSearch.aspx', fourth),
        # #             ('http://bankrot.fedresurs.ru/DisqualificantsList.aspx', fifth),
        # #             ('https://service.nalog.ru/uwsfind.do', sixth if sixth else 'ошибка ввода'),
        # #             ('https://service.nalog.ru/disqualified.do', seventh)
        # #     )

        self.info = dict_div
        driver.close()
        return self.fio


if __name__ == '__main__':
    services = (
                'http://zakupki.gov.ru/epz/dishonestsupplier/quicksearch/search.html',
                'https://service.nalog.ru/bi.do',
                'http://se.fedresurs.ru/IndividualEntrepreneurs',
        )
    headers = (
        'Портал Закупок Результаты поиска',
        'Система информирования банков о состоянии обработки электронных документов (311-П, 440-П)',
        'Единый федеральный реестр',

        # 'Поиск сведений в реестре дисквалифицированных лиц',
        # 'Единый федеральный реестр сведений о банкротстве',
        # 'Сведения о юридических лицах и индивидуальных предпринимателях, в отношении которых представлены документы для государственной регистрации',
        # 'Поиск сведений в реестре дисквалифицированных лиц',
    )



    key = ''
    ogrn = ''

    file = ('{} - {}{}').format(key, dt.today().strftime("%d-%m-%Y"), '.docx')

    key = [x for x in key]
    ogrn = [x for x in ogrn]

    fl = Fl((key, ogrn),)
    fl.services()

    dict_div = fl.info

    document = Document()
    index_header = 0
    for k, v in dict_div.items():
        soup = bs(v, 'lxml')


        head = headers[index_header]
        index_header += 1

        document.add_heading(head, level=2)
        try:
            # soup_table = soup.find('table', class_='search-result')
            soup_table = soup.find('table')
            print('soup_table is ok!')
        except Exception as e:
            print('soup_table is not ok!' + e)

        if soup_table:
            soup_rows = soup_table.find_all('tr')
            try:
                soup_cols = soup_table.find_all('tr')[1].find_all('td')
            except IndexError:
                soup_cols = soup_table.find_all('tr')[0].find_all('td')

            len_rows = len(soup_rows)
            len_cols = len(soup_cols)

            doc_table = document.add_table(rows=len_rows, cols=len_cols, style='Table Grid')

            for idx, soup_row in enumerate(soup_rows):
                soup_cols = soup_row.find_all('td') or soup_row.find_all('th')

                doc_cells = doc_table.rows[idx].cells
                for i, soup_col in enumerate(soup_cols):
                    doc_cells[i].text = soup_col.text
        else:
            try:
                nf = soup.find('p').text
            except NoSuchElementException:
                nf = soup.get_text()
            except AttributeError:
                nf = soup.get_text()

            document.add_paragraph(nf)
            # print('ой, тут не таблица')
        document.add_paragraph()

    document.save(file)
    print('сохраняем документ')
